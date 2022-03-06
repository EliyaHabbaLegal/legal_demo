import os

import docx
import torch
from torch import nn

from answer import Answer

ALEPHBERT_PARASHOOT = "alephbert-fine-tuning-ParaShoot"
LEGALHBERT_PARASHOOT = "legal-alephbert-fine-tuning-ParaShoot"
ALEPHBERT_SQUAD = "alephbert-fine-tuning-squad"
LEGALHBERT_SQUAD = "legal-alephbert-fine-squad"
MBERT_PARASHOOT = "mbert-fine-tuning-ParaShoot"
MODELS = [ALEPHBERT_PARASHOOT, LEGALHBERT_PARASHOOT, MBERT_PARASHOOT, ALEPHBERT_SQUAD, LEGALHBERT_SQUAD]
models_names_to_models_path = {ALEPHBERT_PARASHOOT: os.path.join('models', 'finetuining_aleph_onparashoot', "epoch_4"),
                               LEGALHBERT_PARASHOOT: os.path.join('models', 'finetuining_legalaleph_onparashoot',
                                                                  "epoch_5"),
                               MBERT_PARASHOOT: os.path.join('models', 'finetuining_mbert_onparashoot', "epoch_10"),
                               LEGALHBERT_SQUAD: os.path.join('models', 'finetuining_legalaleph_squad', "epoch_10"),
                               ALEPHBERT_SQUAD: os.path.join('models', 'finetuining_aleph_squad', "epoch_10"),
                               }

LEGAL_QA_EXAMPLES_PATH = os.path.join('data', 'legal_qa.jsonl')

def read_file_to_pars(uploaded_file):
    doc = docx.Document(uploaded_file)
    paragraphs = doc.paragraphs
    return paragraphs


def get_list_of_epochs_of_model(model):
    model_dir = os.path.dirname(models_names_to_models_path[model])
    epochs = os.listdir(model_dir)

    epochs_paths = {epoch: os.path.join(model_dir, epoch) for epoch in epochs}
    return epochs_paths


def predict(model, tokenizer, encoding, top_k):
    input_ids, attention_mask = encoding["input_ids"], encoding["attention_mask"]
    with torch.no_grad():
        start_scores, end_scores = model(torch.tensor([input_ids]),
                                         attention_mask=torch.tensor([attention_mask])).values()
    softmax = nn.Softmax(dim=0)
    start_scores_probs = softmax(torch.squeeze(start_scores))
    end_scores_probs = softmax(torch.squeeze(end_scores))

    values_start, indices_start = torch.topk(start_scores_probs, top_k)
    values_end, indices_end = torch.topk(end_scores_probs, top_k)
    answers = []
    for ans_i in range(top_k):
        answer_start_ind, answer_end_index = indices_start[ans_i].item(), indices_end[ans_i].item()
        answer_start_confidence, answer_end_confidence = values_start[ans_i].item(), values_end[ans_i].item()
        answer = Answer(tokenizer, input_ids, answer_start_ind, answer_end_index, answer_start_confidence,
                        answer_end_confidence)
        answer.combine_answer_from_context()
        answers.append(answer)
    return answers



def save_results_to_docx(q, answers, confidences, model_name):
    doc = docx.Document()
    p_conf = doc.add_paragraph()
    p_conf.add_run("question: ")
    p_conf.add_run(q)

    for answer, confidence in zip(answers, confidences):
        p_conf = doc.add_paragraph()
        p_conf.add_run("confidence: ")
        p_conf.add_run(confidence)
        # red = RGBColor(255, 0, 0)
        # p_conf.font.color.rgb = red

        p = doc.add_paragraph()
        prefix, real_answer, suffix = answer.beginning_of_context_until_answer, answer.answer_itself, answer.end_of_context_from_the_answer
        p.add_run(prefix)
        p.add_run(" ")
        p.add_run(real_answer).bold = True
        p.add_run(" ")
        p.add_run(suffix)
    output_path = os.path.join("results_from_demo", model_name + '.docx')
    doc.save(output_path)

    return output_path
